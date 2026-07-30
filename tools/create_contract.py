from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documenti"
OUT.mkdir(exist_ok=True)
OUTPUT = OUT / "Contratto_realizzazione_sito_web_La_Speziale.docx"

INK = "26332C"
SAGE = "5D785F"
LIGHT = "E8EEE5"
MUTED = "667168"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", *, bold_prefix=None, italic=False, after=6, align=None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    p.paragraph_format.keep_together = keep
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text))
    return p


def add_clause(doc, number, title, paragraphs=None, bullets=None):
    h = doc.add_paragraph()
    h.style = doc.styles["Heading 1"]
    h.paragraph_format.keep_with_next = True
    set_run(h.add_run(f"Art. {number} - {title}"), size=12.5, bold=True, color=SAGE)
    for text in paragraphs or []:
        add_para(doc, text)
    for text in bullets or []:
        add_bullet(doc, text)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.1)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.35)
section.right_margin = Cm(2.35)
section.header_distance = Cm(0.9)
section.footer_distance = Cm(0.9)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.16

for style_name in ("Heading 1", "Heading 2"):
    style = styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.color.rgb = RGBColor.from_string(SAGE)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(hp.add_run("LA SPEZIALE | ACCORDO PER LA REALIZZAZIONE DEL SITO WEB"), size=8, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(fp.add_run("Contratto di prestazione d'opera occasionale - 24 giugno 2026"), size=8, color=MUTED)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(16)
title.paragraph_format.space_after = Pt(8)
set_run(title.add_run("CONTRATTO DI PRESTAZIONE D'OPERA OCCASIONALE"), size=20, bold=True, color=INK)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(22)
set_run(subtitle.add_run("Realizzazione del sito web e della piattaforma di gestione dei contenuti"), size=12, bold=True, color=SAGE)

party_table = doc.add_table(rows=1, cols=2)
party_table.alignment = WD_TABLE_ALIGNMENT.CENTER
party_table.autofit = False
party_table.columns[0].width = Cm(7.75)
party_table.columns[1].width = Cm(7.75)
left, right = party_table.rows[0].cells
for cell in (left, right):
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

lp = left.paragraphs[0]
set_run(lp.add_run("CLIENTE"), size=8.5, bold=True, color=SAGE)
for text, bold in [
    ("LA SPEZIALE S.R.L.", True),
    ("P. IVA e C.F. 12756050154", False),
    ("Corso Giacomo Matteotti n. 3", False),
    ("20121 Milano (MI)", False),
    ("PEC: laspeziale@lamiapec.it", False),
]:
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(text), bold=bold)

rp = right.paragraphs[0]
set_run(rp.add_run("FORNITORE"), size=8.5, bold=True, color=SAGE)
for text, bold in [
    ("Davide Doneddu", True),
    ("C.F. DNDDVD02R01F704H", False),
    ("Residenza: [INSERIRE INDIRIZZO]", False),
    ("Email/PEC: [INSERIRE]", False),
]:
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(text), bold=bold)

add_para(doc, "", after=8)
add_para(
    doc,
    'Il Cliente e il Fornitore sono di seguito indicati singolarmente come la "Parte" e congiuntamente come le "Parti".',
    italic=True,
    after=10,
)

add_clause(
    doc,
    1,
    "Premesse e natura del rapporto",
    [
        "Le premesse e gli eventuali allegati costituiscono parte integrante del presente contratto.",
        "Il Cliente intende rinnovare la propria presenza online mediante un nuovo sito web e un sistema di gestione dei contenuti. Il Fornitore dichiara di possedere le competenze necessarie per eseguire l'incarico con lavoro prevalentemente proprio, autonomia organizzativa e senza vincolo di subordinazione, ai sensi degli articoli 2222 e seguenti del Codice civile.",
        "Il rapporto ha natura occasionale e non comporta inserimento del Fornitore nell'organizzazione aziendale del Cliente, obblighi di presenza, esclusiva o continuità oltre quanto espressamente previsto.",
    ],
)

add_clause(
    doc,
    2,
    "Oggetto dell'incarico",
    [
        "Il Fornitore si obbliga a progettare, sviluppare, configurare e rendere disponibile per la pubblicazione un sito web personalizzato per La Speziale, comprensivo di un'area di gestione dei contenuti basata su Sanity Studio.",
        "Il progetto comprende la realizzazione e configurazione delle seguenti aree:",
    ],
    [
        "pagine istituzionali: Homepage, Chi siamo, Dicono di noi, Contatti, Privacy Policy e Cookie Policy;",
        "sezione Programmi, con pagina generale, pagine di dettaglio e gestione dei contenuti da Studio;",
        "sezione Prodotti, con pagina generale, ricerca e filtri, pagine di dettaglio, collegamenti esterni e gestione dei contenuti da Studio;",
        "sezione Fitoterapia, con pagina generale, pagine di dettaglio, collegamenti con i prodotti consigliati e gestione dei contenuti da Studio;",
        "navigazione responsive, footer, metadati SEO di base, gestione immagini, call to action e contenuti globali;",
        "configurazione degli schemi Sanity, delle anteprime editoriali e dei permessi tecnici necessari al normale utilizzo.",
    ],
)

add_clause(
    doc,
    3,
    "Deliverable e attività incluse",
    [
        "Sono inclusi nel corrispettivo:",
    ],
    [
        "progettazione grafica e sviluppo responsive per desktop, tablet e smartphone;",
        "implementazione delle funzionalità descritte all'articolo 2;",
        "configurazione iniziale del progetto Astro e del Sanity Studio;",
        "migrazione o inserimento dei contenuti concordati e disponibili al momento dello sviluppo;",
        "configurazione tecnica SEO di base, incluse struttura semantica, meta title, meta description, URL e testi alternativi per le immagini quando forniti;",
        "test funzionali sui browser moderni più diffusi;",
        "breve guida operativa e sessione di consegna per l'utilizzo dello Studio;",
        "assistenza post-consegna nei limiti dell'articolo 12.",
    ],
)

add_clause(
    doc,
    4,
    "Attività escluse",
    [
        "Salvo diverso accordo scritto, non sono inclusi:",
    ],
    [
        "registrazione o rinnovo del dominio, hosting, caselle email e servizi a pagamento di terzi;",
        "redazione professionale di testi, servizi fotografici, traduzioni o produzione di contenuti non già forniti;",
        "consulenza legale, fiscale, medica o specialistica sui contenuti del sito;",
        "e-commerce, pagamenti online, autenticazione utenti, aree riservate o integrazioni non indicate all'articolo 2;",
        "manutenzione evolutiva, nuove pagine, nuove funzionalità o redesign successivi all'accettazione;",
        "attività continuativa di caricamento o aggiornamento dei contenuti per conto del Cliente.",
    ],
)

add_clause(
    doc,
    5,
    "Collaborazione e materiali del Cliente",
    [
        "Il Cliente si impegna a nominare un referente, fornire tempestivamente testi, immagini, loghi, credenziali e informazioni necessarie, nonché a formulare approvazioni e osservazioni in modo unitario.",
        "I termini di consegna restano sospesi per il tempo in cui il Fornitore non possa procedere a causa della mancata consegna dei materiali, di risposte incomplete o di ritardi nelle approvazioni.",
        "Il Fornitore può utilizzare contenuti provvisori o segnaposto per proseguire lo sviluppo. La loro sostituzione successiva non costituisce modifica progettuale, purché non richieda variazioni strutturali.",
    ],
)

add_clause(
    doc,
    6,
    "Tempistiche",
    [
        "Il primo rilascio funzionante sarà reso disponibile entro 7 giorni lavorativi dalla data più recente tra: sottoscrizione del contratto, ricezione dell'acconto, approvazione dello stile grafico e consegna dei materiali indispensabili.",
        "Il termine è indicativo e sarà ragionevolmente prorogato in caso di richieste aggiuntive, variazioni dei requisiti, ritardi del Cliente, indisponibilità di servizi terzi, forza maggiore o circostanze tecniche non prevedibili con l'ordinaria diligenza.",
        "Le Parti potranno concordare per iscritto scadenze intermedie o priorità diverse.",
    ],
)

add_clause(
    doc,
    7,
    "Modifiche e attività extra",
    [
        "Il corrispettivo include un ciclo organico di revisioni sul rilascio presentato, purché le richieste siano coerenti con l'oggetto originario e comunicate entro il periodo di collaudo.",
        "Sono considerate attività extra le variazioni strutturali, le nuove funzionalità, l'aggiunta di pagine o sezioni non previste, il rifacimento di parti già approvate e ogni intervento determinato da nuovi requisiti.",
        "Le attività extra saranno eseguite soltanto dopo approvazione scritta di tempi e corrispettivo, anche tramite email o PEC.",
    ],
)

add_clause(
    doc,
    8,
    "Collaudo e accettazione",
    [
        "Il Fornitore comunicherà al Cliente la disponibilità del progetto per il collaudo. Il Cliente avrà 7 giorni lavorativi per verificare il sito e segnalare per iscritto eventuali difformità rispetto agli articoli 2 e 3.",
        "Il Fornitore correggerà senza costi le anomalie riproducibili e imputabili al proprio lavoro. Non costituiscono anomalie le preferenze estetiche sopravvenute, le modifiche richieste dopo approvazione, i problemi causati da servizi terzi o gli interventi eseguiti dal Cliente o da altri soggetti.",
        "Il progetto si considera accettato alla prima tra le seguenti circostanze: conferma scritta del Cliente, pubblicazione sul dominio definitivo, utilizzo operativo del sito oppure decorso del termine di collaudo senza contestazioni specifiche.",
    ],
)

add_clause(
    doc,
    9,
    "Corrispettivo e trattamento fiscale",
    [
        "Il compenso complessivo è stabilito in euro 1.600,00 lordi (milleseicento/00). L'importo è riferito alla prestazione e sarà assoggettato alla ritenuta d'acconto prevista dalla normativa applicabile, ove dovuta.",
        "Il Fornitore emetterà ricevuta per prestazione occasionale. L'eventuale imposta di bollo e gli eventuali obblighi contributivi saranno gestiti secondo la normativa vigente e le rispettive responsabilità delle Parti.",
        "Il pagamento è così ripartito:",
    ],
    [
        "acconto del 40%, pari a euro 640,00 lordi, alla sottoscrizione del contratto;",
        "saldo del 60%, pari a euro 960,00 lordi, al completamento del collaudo e comunque prima della pubblicazione definitiva o della consegna finale.",
    ],
)

add_clause(
    doc,
    10,
    "Modalità e ritardi di pagamento",
    [
        "I pagamenti saranno effettuati tramite bonifico bancario intestato a Davide Doneddu, IBAN [INSERIRE IBAN].",
        'Causale acconto: "Acconto realizzazione sito web La Speziale S.r.l.". Causale saldo: "Saldo realizzazione sito web La Speziale S.r.l.".',
        "In caso di ritardo superiore a 7 giorni, previa comunicazione scritta, il Fornitore potrà sospendere le attività, la pubblicazione, l'assistenza o la consegna delle componenti ancora sotto il proprio controllo fino all'avvenuto pagamento. Restano applicabili gli interessi di mora previsti dalla legge.",
    ],
)

add_clause(
    doc,
    11,
    "Dominio, hosting, servizi esterni e credenziali",
    [
        "Dominio, hosting, caselle email, account Sanity e ogni servizio esterno dovranno essere intestati al Cliente oppure trasferiti a esso quando tecnicamente possibile. I relativi canoni, rinnovi e condizioni economiche sono a carico del Cliente.",
        "Le credenziali di amministrazione e i contenuti inseriti negli account del Cliente appartengono al Cliente. Il Cliente è responsabile della loro conservazione, della gestione degli accessi e dell'immediata revoca delle autorizzazioni non più necessarie.",
        "Il Fornitore non risponde di interruzioni, perdite di dati, modifiche dei prezzi, variazioni contrattuali, dismissioni o malfunzionamenti imputabili ai fornitori terzi.",
    ],
)

add_clause(
    doc,
    12,
    "Assistenza post-consegna",
    [
        "Dalla data di accettazione decorre un periodo di assistenza incluso della durata di 4 mesi.",
        "L'assistenza comprende chiarimenti sull'uso di Sanity Studio e la correzione di anomalie del codice originariamente consegnato e riproducibili nell'ambiente concordato.",
        "Sono esclusi dall'assistenza: inserimento contenuti, modifiche grafiche, nuove funzionalità, aggiornamenti richiesti da servizi terzi, ripristino di dati cancellati, problemi derivanti da credenziali compromesse, modifiche effettuate da terzi e interventi su infrastrutture non gestite dal Fornitore.",
        "Gli interventi esclusi potranno essere oggetto di separato preventivo. L'assistenza non equivale a un servizio continuativo di manutenzione, monitoraggio, backup o presidio di sicurezza.",
    ],
)

add_clause(
    doc,
    13,
    "Proprietà intellettuale e licenza d'uso",
    [
        "A seguito del pagamento integrale, il Cliente acquisisce una licenza d'uso perpetua, non esclusiva e senza limiti territoriali del sito realizzato, finalizzata alla pubblicazione, gestione e promozione della propria attività.",
        "Restano di proprietà del Fornitore il codice sorgente, l'architettura, gli strumenti, i metodi, i moduli generici, i componenti riutilizzabili e il know-how sviluppati prima o durante l'incarico. Il Fornitore può riutilizzare tali elementi in altri progetti, senza divulgare contenuti riservati o dati del Cliente.",
        "Il presente corrispettivo comprende la consegna o messa a disposizione degli artefatti necessari alla pubblicazione e alla gestione ordinaria concordata, ma non la cessione esclusiva del codice sorgente. L'eventuale consegna del repository sorgente, il trasferimento tecnico o la cessione dei diritti patrimoniali dovranno essere concordati separatamente per iscritto.",
        "Restano ferme le licenze dei software e delle librerie di terzi, che continuano a essere disciplinate dalle rispettive condizioni.",
    ],
)

add_clause(
    doc,
    14,
    "Contenuti, conformità e responsabilità editoriale",
    [
        "Il Cliente è l'unico responsabile dell'esattezza, liceità, aggiornamento e adeguatezza di testi, immagini, marchi, recensioni, indicazioni relative a prodotti, fitoterapia, nutrizione e attività professionali pubblicate sul sito.",
        "Il Cliente garantisce di possedere i diritti e le autorizzazioni necessari sui materiali forniti e terrà indenne il Fornitore da pretese di terzi derivanti dal loro utilizzo.",
        "Privacy Policy, Cookie Policy, banner e strumenti tecnici eventualmente predisposti dal Fornitore costituiscono implementazione tecnica sulla base delle informazioni ricevute e non sostituiscono una consulenza legale. Il Cliente resta responsabile della verifica e dell'approvazione dei testi legali e della configurazione dei servizi di tracciamento effettivamente utilizzati.",
    ],
)

add_clause(
    doc,
    15,
    "Protezione dei dati personali",
    [
        "Le Parti trattano i dati personali acquisiti per l'esecuzione del contratto nel rispetto del Regolamento (UE) 2016/679 e della normativa nazionale applicabile, ciascuna secondo il proprio ruolo.",
        "Qualora il Fornitore debba trattare dati personali per conto del Cliente in modo non meramente occasionale o accedere a dati degli utenti del sito, le Parti sottoscriveranno, prima del trattamento, un separato accordo ai sensi dell'articolo 28 del Regolamento (UE) 2016/679.",
        "Il sito non comprenderà raccolte di dati o integrazioni ulteriori rispetto a quelle espressamente concordate.",
    ],
)

add_clause(
    doc,
    16,
    "Riservatezza",
    [
        "Ciascuna Parte si impegna a non divulgare informazioni tecniche, commerciali, credenziali, documenti e dati dell'altra Parte appresi durante l'esecuzione dell'incarico, salvo quanto necessario per l'esecuzione del contratto o richiesto dalla legge.",
        "L'obbligo di riservatezza resta efficace per 3 anni dalla cessazione del rapporto e non si applica alle informazioni già pubbliche, legittimamente conosciute o sviluppate in modo indipendente.",
    ],
)

add_clause(
    doc,
    17,
    "Garanzie e limitazione di responsabilità",
    [
        "Il Fornitore eseguirà l'incarico con diligenza professionale e secondo le specifiche concordate, ma non garantisce risultati economici, commerciali, di posizionamento sui motori di ricerca o continuità assoluta dei servizi esterni.",
        "Fatti salvi dolo, colpa grave e responsabilità non limitabili per legge, il Fornitore non risponde di danni indiretti, perdita di profitti, perdita di opportunità, danni reputazionali, indisponibilità di servizi terzi, attacchi a infrastrutture non gestite dal Fornitore, uso improprio dello Studio o modifiche effettuate da soggetti diversi dal Fornitore.",
        "Nei limiti consentiti dalla legge, la responsabilità complessiva del Fornitore connessa al presente contratto non potrà superare il corrispettivo effettivamente pagato dal Cliente.",
    ],
)

add_clause(
    doc,
    18,
    "Recesso, risoluzione e cessazione",
    [
        "Il Cliente può recedere in qualsiasi momento mediante comunicazione scritta. In tal caso l'acconto resta acquisito a compensazione delle attività avviate; se il valore del lavoro eseguito e degli impegni assunti supera l'acconto, il Cliente corrisponderà la differenza documentata in proporzione allo stato di avanzamento.",
        "Il Fornitore può recedere per giusta causa o risolvere il contratto in caso di mancato pagamento, mancata collaborazione protratta per oltre 15 giorni, richieste illecite o comportamento che renda irragionevolmente impossibile la prosecuzione, previa comunicazione scritta e termine di 7 giorni per porre rimedio, quando applicabile.",
        "In caso di cessazione, il Fornitore consegnerà quanto già pagato e tecnicamente utilizzabile, fermo restando l'articolo 13.",
    ],
)

add_clause(
    doc,
    19,
    "Forza maggiore",
    [
        "Nessuna Parte sarà responsabile per ritardi o inadempimenti dovuti a eventi fuori dal ragionevole controllo, inclusi guasti generalizzati, indisponibilità di piattaforme, provvedimenti dell'autorità, eventi naturali, malattia grave o impedimenti tecnici non prevedibili. La Parte interessata informerà tempestivamente l'altra e le scadenze saranno prorogate per la durata dell'impedimento.",
    ],
)

add_clause(
    doc,
    20,
    "Portfolio e riferimenti professionali",
    [
        "Dopo la pubblicazione del sito, il Cliente autorizza il Fornitore a indicare La Speziale tra i propri clienti e a mostrare immagini pubbliche del progetto nel portfolio, sul sito professionale e sui canali di comunicazione del Fornitore.",
        "Sono esclusi dati riservati, credenziali, statistiche non pubbliche e informazioni interne. Il Cliente può revocare l'autorizzazione per comprovate esigenze di riservatezza mediante comunicazione scritta.",
    ],
)

add_clause(
    doc,
    21,
    "Comunicazioni",
    [
        "Le comunicazioni operative possono avvenire tramite email o altri canali concordati. Le comunicazioni relative a recesso, risoluzione, contestazioni formali e variazioni economiche devono essere inviate tramite email con conferma di ricezione o PEC agli indirizzi indicati nel presente contratto.",
        "Ciascuna Parte comunicherà tempestivamente eventuali variazioni dei propri recapiti.",
    ],
)

add_clause(
    doc,
    22,
    "Legge applicabile e foro competente",
    [
        "Il contratto è regolato dalla legge italiana. Le Parti si impegnano a tentare una soluzione amichevole delle controversie entro 30 giorni dalla contestazione scritta.",
        "Per ogni controversia relativa a validità, interpretazione, esecuzione o cessazione del contratto sarà competente in via esclusiva il Foro di Milano, salvo diversa competenza inderogabile prevista dalla legge.",
    ],
)

add_clause(
    doc,
    23,
    "Disposizioni finali",
    [
        "Il presente contratto sostituisce ogni precedente intesa sul medesimo oggetto. Modifiche e integrazioni sono valide soltanto se formulate per iscritto e accettate da entrambe le Parti.",
        "L'eventuale nullità o inefficacia di una clausola non pregiudica la validità delle restanti disposizioni. La mancata applicazione occasionale di una clausola non costituisce rinuncia.",
        "Il contratto può essere sottoscritto con firma autografa o elettronica e scambiato in copie, ciascuna delle quali sarà considerata originale.",
    ],
)

doc.add_page_break()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.paragraph_format.space_after = Pt(18)
set_run(closing.add_run("SOTTOSCRIZIONE"), size=16, bold=True, color=SAGE)

add_para(doc, "Letto, confermato e sottoscritto.", align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
add_para(doc, "Milano, 24 giugno 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=26)

sign_table = doc.add_table(rows=1, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sign_table.autofit = False
sign_table.columns[0].width = Cm(7.75)
sign_table.columns[1].width = Cm(7.75)
for cell in sign_table.rows[0].cells:
    set_cell_margins(cell, top=100, bottom=100, start=120, end=120)

for cell, title_text, lines in [
    (sign_table.cell(0, 0), "IL CLIENTE", ["La Speziale S.r.l.", "Nome e qualifica: ____________________", "", "Firma: ______________________________"]),
    (sign_table.cell(0, 1), "IL FORNITORE", ["Davide Doneddu", "", "", "Firma: ______________________________"]),
]:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(title_text), size=9, bold=True, color=SAGE)
    for line in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        set_run(p.add_run(line), size=10, bold=line in ("La Speziale S.r.l.", "Davide Doneddu"))

add_para(doc, "", after=10)
specific = doc.add_paragraph()
specific.paragraph_format.space_after = Pt(8)
set_run(specific.add_run("APPROVAZIONE SPECIFICA DELLE CLAUSOLE"), size=11.5, bold=True, color=SAGE)
add_para(
    doc,
    "Ai sensi e per gli effetti degli articoli 1341 e 1342 del Codice civile, il Cliente dichiara di approvare specificamente le clausole di cui agli articoli: 6 (Tempistiche), 7 (Modifiche e attività extra), 8 (Collaudo e accettazione), 10 (Ritardi e sospensione), 11 (Servizi esterni), 12 (Limiti dell'assistenza), 13 (Proprietà intellettuale e licenza), 14 (Responsabilità editoriale), 17 (Limitazione di responsabilità), 18 (Recesso e risoluzione), 20 (Portfolio) e 22 (Foro competente).",
    after=18,
)
add_para(doc, "Per LA SPEZIALE S.R.L.", bold_prefix="Per LA SPEZIALE S.R.L.", after=10)
add_para(doc, "Firma: ______________________________________________", after=18)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(18)
note.paragraph_format.space_after = Pt(0)
set_run(
    note.add_run(
        "Nota operativa: completare i campi tra parentesi quadre e far verificare il testo da un professionista legale o fiscale prima della firma."
    ),
    size=8.5,
    italic=True,
    color=MUTED,
)

doc.core_properties.title = "Contratto realizzazione sito web La Speziale"
doc.core_properties.subject = "Prestazione d'opera occasionale per sito web e Sanity Studio"
doc.core_properties.author = "Davide Doneddu"
doc.core_properties.keywords = "contratto, sito web, Astro, Sanity, La Speziale"

doc.save(OUTPUT)
print(OUTPUT)
