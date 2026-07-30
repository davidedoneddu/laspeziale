# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\ACER\source\repos\laspeziale")
OUT = ROOT / "output" / "docs" / "Guida-operativa-La-Speziale-editabile.docx"
LOGO = ROOT / "apps" / "web" / "public" / "assets" / "old-site" / "logo.png"
HERO = ROOT / "apps" / "web" / "public" / "assets" / "old-site" / "hero-header.jpg"
NUTRITION = ROOT / "apps" / "web" / "public" / "assets" / "old-site" / "nutrizione-integrata.jpg"
FITOTERAPIA = ROOT / "apps" / "web" / "public" / "assets" / "old-site" / "fitoterapia.jpg"

INK = "26332C"
MUTED = "647066"
SAGE = "5D785F"
SAGE_MID = "8FA58F"
SAGE_PALE = "E7EBE2"
IVORY = "FBF7EE"
WHITE = "FFFDF8"
BEIGE = "E8DED0"
GOLD = "A07A42"
RED = "9B3A3A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_alt_text(shape, description):
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_picture(doc, path, width, alt, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5 if caption else 10)
    shape = p.add_run().add_picture(str(path), width=width)
    add_alt_text(shape, alt)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        set_run(cp.add_run(caption), 8.5, MUTED, italic=True)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run(run, 8.5, MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, before, after, color in (
        ("Heading 1", 18, 18, 9, SAGE),
        ("Heading 2", 14, 14, 7, SAGE),
        ("Heading 3", 11.5, 10, 4, INK),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def configure_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78 if first else 0.85)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.35)
    if first:
        section.different_first_page_header_footer = True
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("LA SPEZIALE  |  GUIDA OPERATIVA"), 8, SAGE, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.add_run("La Speziale  |  Gestione sito e contenuti")
    set_run(fp.runs[0], 8, MUTED)
    add_page_number(fp)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(text.upper()), 8.5, SAGE, bold=True)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), 10.5, INK, bold=True)
        set_run(p.add_run(text[len(bold_prefix):]), 10.5, MUTED)
    else:
        set_run(p.add_run(text), 10.5, MUTED)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), 10.5, INK, bold=True)
        set_run(p.add_run(text[len(bold_prefix):]), 10.5, MUTED)
    else:
        set_run(p.add_run(text), 10.5, MUTED)
    return p


def add_number(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), 10.5, INK, bold=True)
        set_run(p.add_run(text[len(bold_prefix):]), 10.5, MUTED)
    else:
        set_run(p.add_run(text), 10.5, MUTED)
    return p


def add_callout(doc, title, text, kind="info"):
    palette = {
        "info": (SAGE_PALE, SAGE),
        "tip": (IVORY, GOLD),
        "warning": ("F8EAEA", RED),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), 10.5, accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(text), 9.7, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_two_col_table(doc, rows, widths=(2500, 6860), header=None):
    table = doc.add_table(rows=0, cols=2)
    if header:
        row = table.add_row()
        for idx, value in enumerate(header):
            set_cell_shading(row.cells[idx], SAGE)
            set_run(row.cells[idx].paragraphs[0].add_run(value), 9.2, WHITE, bold=True)
        set_repeat_table_header(row)
    for label, detail in rows:
        row = table.add_row()
        set_cell_shading(row.cells[0], SAGE_PALE)
        set_run(row.cells[0].paragraphs[0].add_run(label), 9.4, INK, bold=True)
        set_run(row.cells[1].paragraphs[0].add_run(detail), 9.4, MUTED)
    set_table_geometry(table, list(widths))
    table.style = "Table Grid"
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_chapter(doc, number, title, subtitle=None):
    doc.add_page_break()
    add_kicker(doc, f"Capitolo {number}")
    h = doc.add_paragraph(style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    set_run(h.add_run(title), 22, SAGE, bold=True, font="Aptos Display")
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        set_run(p.add_run(subtitle), 11.5, MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "Guida operativa La Speziale"
    doc.core_properties.subject = "Gestione del sito web e di Sanity Studio"
    doc.core_properties.author = "La Speziale"
    doc.core_properties.keywords = "La Speziale, Sanity Studio, guida, sito web"
    configure_styles(doc)
    configure_section(doc.sections[0], first=True)

    # Cover: editorial_cover with La Speziale brand override.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(2.35))
    add_alt_text(logo_shape, "Logo La Speziale")
    add_kicker(doc, "Manuale per la gestione dei contenuti")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("Guida operativa\ndel sito La Speziale"), 29, SAGE, bold=True, font="Aptos Display")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(18)
    set_run(sub.add_run("Come aggiornare testi, immagini, programmi e prodotti in autonomia, mantenendo il layout protetto."), 13, MUTED)
    add_picture(doc, HERO, Inches(6.5), "Immagine astratta della testata La Speziale")
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(10)
    meta.paragraph_format.space_after = Pt(0)
    set_run(meta.add_run("Versione 1.0  |  Giugno 2026"), 9.5, MUTED, bold=True)
    add_callout(doc, "Documento modificabile", "Il file Word allegato è il sorgente del manuale. Può essere aggiornato e riconvertito in PDF quando cambiano funzioni o contenuti del sito.", "tip")

    add_chapter(doc, "1", "Il sistema in breve", "Due strumenti distinti, collegati tra loro: il sito pubblico e lo Studio dei contenuti.")
    add_two_col_table(doc, [
        ("Sito pubblico", "È ciò che vedono visitatori e clienti. Mostra pagine, programmi, prodotti, contatti, recensioni e informative."),
        ("Sanity Studio", "È il pannello riservato in cui si aggiornano testi, immagini e contenuti senza intervenire sul codice."),
        ("Layout protetto", "Grafica, spaziature, componenti e comportamento responsive sono gestiti dal sito. Lo Studio modifica solo i contenuti previsti."),
        ("Pubblicazione", "Una modifica diventa pubblica solo quando il documento viene pubblicato. Le bozze possono essere controllate in anteprima."),
    ], header=("Elemento", "A cosa serve"))
    add_callout(doc, "La regola più importante", "Modifica i campi disponibili, controlla l'anteprima e pubblica. Non servono competenze tecniche e non è possibile rompere il layout usando i campi predisposti.")
    doc.add_heading("Cosa viene aggiornato automaticamente", level=2)
    add_bullet(doc, "Un programma pubblicato compare nella pagina Programmi e, se abilitato, nel menu.")
    add_bullet(doc, "Un prodotto pubblicato compare nella griglia Prodotti e nella ricerca testuale.")
    add_bullet(doc, "Contatti, CTA globali e dati legali vengono riutilizzati nelle pagine che ne hanno bisogno.")
    add_bullet(doc, "Un contenuto disattivato o eliminato non deve più apparire sul sito pubblico.")

    add_chapter(doc, "2", "Accesso e orientamento", "Come entrare nello Studio e riconoscere le tre aree principali.")
    doc.add_heading("Primo accesso", level=2)
    add_number(doc, "Apri l'indirizzo dello Studio fornito dal gestore del sito.")
    add_number(doc, "Accedi con l'account autorizzato, usando Google, GitHub oppure email e password Sanity.")
    add_number(doc, "Attendi il caricamento del menu La Speziale - Gestione sito.")
    add_callout(doc, "Accesso riservato", "Non condividere password o codici di accesso. Ogni collaboratore dovrebbe utilizzare il proprio account Sanity.", "warning")
    doc.add_heading("Il menu dello Studio", level=2)
    add_two_col_table(doc, [
        ("Pagine Statiche", "Impostazioni sito, Homepage, Dicono di noi e Contatti."),
        ("Programmi", "Impostazioni generali della pagina e singoli programmi."),
        ("Prodotti", "Impostazioni generali della pagina e singoli prodotti."),
    ], header=("Area", "Contenuti gestiti"))
    add_body(doc, "Le pagine principali del menu pubblico sono intenzionalmente fisse: Programmi, Prodotti, Contatti e Dicono di noi. Questo protegge navigazione e indicizzazione.")

    add_chapter(doc, "3", "Bozze, anteprima e pubblicazione", "Un flusso semplice per controllare ogni modifica prima che sia visibile.")
    doc.add_heading("Flusso consigliato", level=2)
    add_number(doc, "Apri il documento da modificare.")
    add_number(doc, "Aggiorna testo, immagine o impostazione desiderata.")
    add_number(doc, "Apri la scheda Anteprima per controllare la composizione.")
    add_number(doc, "Correggi eventuali testi troppo lunghi o immagini poco leggibili.")
    add_number(doc, "Premi Pubblica solo quando il contenuto è pronto.")
    add_callout(doc, "Bozza non significa contenuto nascosto", "Se stai modificando un documento già pubblicato, il sito continua a mostrare l'ultima versione pubblicata finché non confermi la nuova pubblicazione.", "tip")
    doc.add_heading("Modifica e Anteprima", level=2)
    add_two_col_table(doc, [
        ("Modifica", "Contiene tutti i campi autorizzati per testi, immagini, collegamenti, visibilità e SEO."),
        ("Anteprima", "Mostra subito come verranno organizzati i contenuti, comprese le modifiche ancora in bozza."),
        ("Apri sito pubblico", "Apre la versione effettivamente pubblicata e visibile ai visitatori."),
    ])

    add_chapter(doc, "4", "Pagine statiche", "Contenuti selezionati e sicuri, senza possibilità di aggiungere strutture che rompano il design.")
    doc.add_heading("Impostazioni sito", level=2)
    add_body(doc, "È il punto centrale per i dati usati in più pagine. Una modifica qui può riflettersi in testata, footer, contatti, CTA e informative.")
    add_two_col_table(doc, [
        ("Identità sito", "Nome, logo, favicon, testo del footer e bottone globale."),
        ("Contatti", "Email, telefono, indirizzo, orari, WhatsApp e social."),
        ("SEO Google", "Titolo e descrizione generali usati come base dal sito."),
        ("Legale", "Titolare, partita IVA o codice fiscale, email privacy e data di aggiornamento delle informative."),
    ])
    doc.add_heading("Homepage", level=2)
    add_bullet(doc, "Hero iniziale: titolo, sottotitolo, immagine e bottone principale.")
    add_bullet(doc, "Sezioni pagina: testo introduttivo, immagini consentite, programmi in evidenza e CTA finale.")
    add_bullet(doc, "SEO Google: titolo, descrizione e immagine di condivisione.")
    doc.add_heading("Dicono di noi", level=2)
    add_bullet(doc, "Impostazioni pagina: apertura, testi, CTA e SEO.")
    add_bullet(doc, "Testimonianze: nome, testo, ruolo, immagine opzionale, visibilità e ordine.")
    add_bullet(doc, "Recensioni Google: sezione separata attivabile, con media, totale, link e selezione delle recensioni mostrate.")
    doc.add_heading("Contatti", level=2)
    add_bullet(doc, "Testi della pagina e del form, titoli dei box e CTA finale.")
    add_bullet(doc, "Posizione della mappa e link Google Maps.")
    add_bullet(doc, "La mappa interattiva viene caricata solo dopo il consenso ai contenuti esterni.")

    add_chapter(doc, "5", "Gestire i programmi", "La pagina generale è separata dai singoli percorsi, che vengono inseriti automaticamente nella griglia.")
    add_picture(doc, NUTRITION, Inches(4.85), "Ingredienti e strumenti per il programma Nutrizione Integrata", "Esempio di immagine programma: orizzontale, chiara e coerente con il contenuto.")
    doc.add_heading("Impostazioni pagina Programmi", level=2)
    add_bullet(doc, "Apertura pagina: etichetta, titolo, testo e immagine hero.")
    add_bullet(doc, "Elenco programmi: titolo, testo e dicitura del bottone delle card.")
    add_bullet(doc, "CTA finale e SEO Google.")
    add_callout(doc, "Griglia automatica", "Non si aggiungono card manualmente nella pagina generale. La griglia prende tutti i programmi pubblicati e li ordina secondo il campo Ordine.")
    doc.add_heading("Creare un nuovo programma", level=2)
    add_number(doc, "Apri Programmi > Elenco programmi e scegli Crea nuovo.")
    add_number(doc, "Compila titolo e indirizzo pagina. Genera lo slug dal titolo e non modificarlo in seguito senza assistenza tecnica.")
    add_number(doc, "Inserisci immagine card, immagine hero, categoria, descrizione breve e testi principali.")
    add_number(doc, "Attiva solo le sezioni realmente utili per quel percorso.")
    add_number(doc, "Completa pubblicazione, ordine, menu e SEO.")
    add_number(doc, "Controlla Anteprima e pubblica.")

    add_chapter(doc, "6", "Comporre una pagina programma", "Ogni sezione può essere accesa o spenta, mentre la grafica resta fissa e coerente.")
    add_two_col_table(doc, [
        ("Menu di ancoraggio", "Collegamenti rapidi alle sezioni attive della pagina."),
        ("Panoramica", "Durata, tipologia, modalità, nota personalizzata e prezzo opzionale."),
        ("A chi si rivolge", "Elenco dei destinatari del percorso."),
        ("Descrizione", "Testo completo che spiega obiettivi e impostazione del percorso."),
        ("Benefici", "Card con icona, titolo e descrizione."),
        ("Come funziona", "Passaggi ordinati del percorso."),
        ("FAQ", "Domande e risposte visualizzate come accordion."),
        ("Contenuti extra", "Sezioni editoriali opzionali predisposte dal sistema."),
        ("CTA finale", "Titolo, testo, immagine e bottoni WhatsApp/email."),
    ], header=("Interruttore", "Effetto sul sito"))
    doc.add_heading("Pubblicazione e menu", level=2)
    add_two_col_table(doc, [
        ("Pubblicato", "Mostra o nasconde il programma dal sito pubblico."),
        ("Visibile nel menu", "Controlla soltanto il sottomenu Programmi."),
        ("In evidenza", "Permette di selezionare il percorso per aree promozionali predisposte."),
        ("Ordine", "Determina la posizione nella griglia e nel menu."),
    ])
    add_callout(doc, "Prima di eliminare", "Se un programma non deve più essere visibile, disattiva Pubblicato. L'eliminazione definitiva va usata solo quando il contenuto non servirà più.", "warning")

    add_chapter(doc, "7", "Gestire i prodotti", "Prodotti ricercabili, pagina dettaglio interna e acquisto su un sito esterno.")
    doc.add_heading("Impostazioni pagina Prodotti", level=2)
    add_bullet(doc, "Hero, introduzione, titoli della griglia e CTA finale.")
    add_bullet(doc, "Etichetta e testo della ricerca, messaggio senza risultati e testo del bottone dettaglio.")
    add_bullet(doc, "La ricerca controlla nome, categoria, descrizione, prezzo, parole chiave e testo della card.")
    doc.add_heading("Creare un prodotto", level=2)
    add_number(doc, "Apri Prodotti > Elenco prodotti e crea un nuovo documento.")
    add_number(doc, "Compila nome e slug, categoria, descrizione breve, prezzo e immagine.")
    add_number(doc, "Inserisci parole chiave utili alla ricerca.")
    add_number(doc, "Aggiungi il link acquisto esterno completo di https:// e il testo del pulsante.")
    add_number(doc, "Compila descrizione completa, caratteristiche e SEO.")
    add_number(doc, "Imposta Pubblicato sul sito e Ordine nella griglia, poi controlla Anteprima.")
    add_two_col_table(doc, [
        ("Scopri il prodotto", "Apre la pagina dettaglio interna del sito La Speziale."),
        ("Acquista", "Apre il sito esterno configurato nel prodotto, in una nuova scheda."),
    ], header=("Pulsante card", "Destinazione"))
    add_callout(doc, "Pulsante acquisto opzionale", "Se il link esterno è vuoto, la card mostra soltanto il collegamento al dettaglio interno. Non inserire link incompleti come www o link senza https://.", "tip")

    add_chapter(doc, "8", "Immagini e testi efficaci", "Poche regole pratiche per mantenere il sito ordinato, credibile e veloce.")
    add_picture(doc, FITOTERAPIA, Inches(5.7), "Prodotti fitoterapici e strumenti da laboratorio", "Un'immagine pertinente aiuta a riconoscere subito il contenuto del programma.")
    doc.add_heading("Immagini", level=2)
    add_bullet(doc, "Usa fotografie nitide, luminose e pertinenti al contenuto.")
    add_bullet(doc, "Preferisci immagini orizzontali per hero e card. Evita screenshot, collage e testi incorporati nelle fotografie.")
    add_bullet(doc, "Compila sempre il testo alternativo descrivendo l'immagine in modo breve e concreto.")
    add_bullet(doc, "Controlla il ritaglio nell'anteprima, soprattutto su smartphone.")
    doc.add_heading("Testi", level=2)
    add_bullet(doc, "Titoli brevi e specifici: una promessa o un argomento per volta.")
    add_bullet(doc, "Descrizioni delle card entro circa 220 caratteri.")
    add_bullet(doc, "Evita maiuscole continue, punti esclamativi ripetuti e informazioni sanitarie non verificate.")
    add_bullet(doc, "Non copiare testi da altri siti. Usa contenuti originali e aggiornati.")

    add_chapter(doc, "9", "SEO e indirizzi delle pagine", "Come mantenere indicizzazione e collegamenti stabili durante gli aggiornamenti.")
    add_two_col_table(doc, [
        ("Titolo SEO", "Descrive la pagina e include il nome La Speziale quando utile. Mantienilo chiaro e non ripetitivo."),
        ("Descrizione SEO", "Riassume il contenuto e invita alla visita con una frase naturale."),
        ("Immagine SEO", "Usata nelle condivisioni social quando disponibile."),
        ("Slug", "È l'indirizzo della pagina. Dopo la pubblicazione non va cambiato senza predisporre un reindirizzamento."),
    ], header=("Campo", "Buona pratica"))
    add_callout(doc, "Protezione dell'indicizzazione", "Le pagine principali e le voci di menu sono fisse. Prima di rinominare o eliminare un programma/prodotto già pubblico, chiedi assistenza per verificare link e reindirizzamenti.", "warning")

    add_chapter(doc, "10", "Privacy, cookie e servizi esterni", "Le funzioni esterne sono integrate con un consenso chiaro e revocabile.")
    add_bullet(doc, "Privacy policy e Cookie policy sono collegate nel footer.")
    add_bullet(doc, "Il banner consente di accettare, rifiutare o personalizzare i contenuti esterni.")
    add_bullet(doc, "Google Maps non viene caricato prima del consenso.")
    add_bullet(doc, "Le preferenze possono essere riaperte dal link Preferenze cookie nel footer.")
    add_bullet(doc, "Dati del titolare, email privacy e data di aggiornamento sono gestiti in Impostazioni sito > Legale.")
    add_callout(doc, "Controllo periodico", "Dati legali, strumenti esterni e informative devono essere verificati ogni volta che cambiano servizi, moduli, analytics o modalità di raccolta dati.", "tip")
    doc.add_heading("Recensioni Google", level=2)
    add_body(doc, "La sezione Recensioni Google è distinta dalle testimonianze interne. Attualmente mostra una selezione inserita nello Studio con valutazione media, numero totale e link alla scheda Google. Non è una sincronizzazione automatica in tempo reale.")

    add_chapter(doc, "11", "Routine consigliata", "Una checklist breve da usare ogni volta che pubblichi un contenuto.")
    doc.add_heading("Prima di pubblicare", level=2)
    for item in (
        "Titolo e descrizione sono chiari e senza refusi.",
        "Immagine e testo alternativo sono presenti.",
        "Link, email, telefono e bottoni portano alla destinazione corretta.",
        "Le sezioni visibili sono solo quelle realmente compilate.",
        "Slug e impostazioni di pubblicazione sono corretti.",
        "Anteprima controllata su desktop e, quando possibile, su smartphone.",
        "Titolo e descrizione SEO sono compilati.",
    ):
        add_bullet(doc, item)
    doc.add_heading("Controllo mensile", level=2)
    add_bullet(doc, "Verifica contatti, orari, prezzi e link esterni.")
    add_bullet(doc, "Rimuovi dalla pubblicazione programmi e prodotti non più disponibili.")
    add_bullet(doc, "Controlla recensioni, testimonianze e contenuti in evidenza.")
    add_bullet(doc, "Aggiorna informative e dati legali quando necessario.")

    add_chapter(doc, "12", "Problemi comuni", "Cosa controllare prima di richiedere assistenza tecnica.")
    add_two_col_table(doc, [
        ("La modifica non appare", "Verifica di aver premuto Pubblica e di stare guardando la pagina corretta. Aggiorna il browser."),
        ("Il programma non è nel menu", "Controlla Pubblicato, Visibile nel menu e Ordine."),
        ("Il prodotto non appare", "Controlla Pubblicato sul sito, slug e immagine obbligatoria."),
        ("Il bottone Acquista manca", "Compila Link acquisto esterno con un URL completo e valido."),
        ("La mappa non si vede", "Accetta i contenuti esterni oppure usa Mostra la mappa. Verifica la posizione in Contatti."),
        ("L'anteprima sembra diversa", "La scheda Anteprima mostra la bozza; il sito pubblico mostra l'ultima versione pubblicata."),
        ("Testo o immagine sono tagliati", "Accorcia il testo o modifica il ritaglio dell'immagine usando hotspot e anteprima."),
    ], header=("Problema", "Controllo rapido"))
    add_callout(doc, "Quando chiedere assistenza", "Contatta il gestore tecnico per nuovi tipi di pagina, modifiche al menu, reindirizzamenti SEO, integrazioni esterne, form realmente funzionanti o cambiamenti al layout.")
    doc.add_heading("Informazioni utili da inviare", level=2)
    add_bullet(doc, "Link della pagina interessata.")
    add_bullet(doc, "Nome del documento aperto nello Studio.")
    add_bullet(doc, "Descrizione breve di ciò che ti aspettavi e di ciò che vedi.")
    add_bullet(doc, "Screenshot dell'errore, senza password, token o dati sensibili.")

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    end_logo = p.add_run().add_picture(str(LOGO), width=Inches(2.2))
    add_alt_text(end_logo, "Logo La Speziale")
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_before = Pt(24)
    set_run(ep.add_run("Contenuti aggiornati, sito protetto."), 22, SAGE, bold=True, font="Aptos Display")
    ep2 = doc.add_paragraph()
    ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(ep2.add_run("Usa questa guida come riferimento e conserva il file Word per gli aggiornamenti futuri."), 11, MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
