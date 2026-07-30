# -*- coding: utf-8 -*-
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape


ROOT = Path(r"C:\Users\ACER\source\repos\laspeziale")
SOURCE = ROOT / "output" / "docs" / "Guida-operativa-La-Speziale-editabile.docx"
OUTPUT = ROOT / "output" / "pdf" / "Guida-operativa-La-Speziale.pdf"

INK = colors.HexColor("#26332C")
MUTED = colors.HexColor("#647066")
SAGE = colors.HexColor("#5D785F")
SAGE_PALE = colors.HexColor("#E7EBE2")
IVORY = colors.HexColor("#FBF7EE")
WHITE = colors.HexColor("#FFFDF8")
GOLD = colors.HexColor("#A07A42")


def iter_blocks(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, parent)


def has_page_break(paragraph):
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def text_with_format(paragraph):
    chunks = []
    for run in paragraph.runs:
        if run._r.xpath('.//w:drawing'):
            continue
        text = escape(run.text).replace("\n", "<br/>")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        chunks.append(text)
    return "".join(chunks) or escape(paragraph.text)


def paragraph_image(paragraph, document):
    blips = paragraph._p.xpath('.//a:blip')
    if not blips:
        return None
    rid = blips[0].get(qn("r:embed"))
    if not rid or rid not in document.part.related_parts:
        return None
    blob = document.part.related_parts[rid].blob
    extents = paragraph._p.xpath('.//wp:extent')
    if extents:
        cx = int(extents[0].get("cx", "0"))
        cy = int(extents[0].get("cy", "0"))
        width = cx / 914400 * inch
        height = cy / 914400 * inch
    else:
        with PILImage.open(BytesIO(blob)) as im:
            width = 6.0 * inch
            height = width * im.height / im.width
    max_width = 6.5 * inch
    max_height = 4.2 * inch
    scale = min(1, max_width / width, max_height / height)
    return Image(BytesIO(blob), width=width * scale, height=height * scale, hAlign="CENTER")


def cell_fill(cell):
    nodes = cell._tc.xpath('./w:tcPr/w:shd')
    if not nodes:
        return None
    value = nodes[0].get(qn("w:fill"))
    if not value or value in ("auto", "FFFFFF"):
        return None
    try:
        return colors.HexColor(f"#{value}")
    except Exception:
        return None


def table_widths(table):
    grid = table._tbl.tblGrid
    widths = [int(col.get(qn("w:w"), "0")) for col in grid.gridCol_lst]
    if not widths or not sum(widths):
        return [6.5 * inch / len(table.columns)] * len(table.columns)
    return [6.5 * inch * value / sum(widths) for value in widths]


def on_page(canvas, doc):
    canvas.saveState()
    if doc.page > 1:
        canvas.setFont("Helvetica-Bold", 7.6)
        canvas.setFillColor(SAGE)
        canvas.drawString(inch, 10.55 * inch, "LA SPEZIALE  |  GUIDA OPERATIVA")
        canvas.setStrokeColor(colors.HexColor("#DDE4DA"))
        canvas.line(inch, 10.38 * inch, 7.5 * inch, 10.38 * inch)
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(MUTED)
    canvas.drawString(inch, 0.42 * inch, "La Speziale  |  Gestione sito e contenuti")
    canvas.drawRightString(7.5 * inch, 0.42 * inch, str(doc.page))
    canvas.restoreState()


def build_pdf():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    source = Document(SOURCE)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.7, leading=12.4, textColor=MUTED, spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=body, fontName="Helvetica-Bold", fontSize=20, leading=23, textColor=SAGE, spaceBefore=2, spaceAfter=11, keepWithNext=True)
    h2 = ParagraphStyle("H2", parent=body, fontName="Helvetica-Bold", fontSize=13.5, leading=16, textColor=SAGE, spaceBefore=12, spaceAfter=7, keepWithNext=True)
    h3 = ParagraphStyle("H3", parent=body, fontName="Helvetica-Bold", fontSize=10.8, leading=13, textColor=INK, spaceBefore=8, spaceAfter=4, keepWithNext=True)
    title_style = ParagraphStyle("Title", parent=body, fontName="Helvetica-Bold", fontSize=27, leading=30, textColor=SAGE, spaceAfter=10)
    subtitle_style = ParagraphStyle("Subtitle", parent=body, fontName="Helvetica", fontSize=12.5, leading=16, textColor=MUTED, spaceAfter=15)
    kicker = ParagraphStyle("Kicker", parent=body, fontName="Helvetica-Bold", fontSize=8.3, leading=10, textColor=SAGE, spaceBefore=2, spaceAfter=6)
    caption = ParagraphStyle("Caption", parent=body, fontName="Helvetica-Oblique", fontSize=8, leading=10, textColor=MUTED, alignment=TA_CENTER, spaceAfter=9)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=18, firstLineIndent=-10, bulletIndent=6, spaceAfter=4)
    number = ParagraphStyle("Number", parent=body, leftIndent=20, firstLineIndent=-14, bulletIndent=4, spaceAfter=4)
    cell_style = ParagraphStyle("Cell", parent=body, fontSize=8.7, leading=11, spaceAfter=0)
    cell_bold = ParagraphStyle("CellBold", parent=cell_style, fontName="Helvetica-Bold", textColor=INK)

    frame = Frame(inch, 0.68 * inch, 6.5 * inch, 9.55 * inch, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    template = PageTemplate(id="guide", frames=[frame], onPage=on_page)
    pdf = BaseDocTemplate(str(OUTPUT), pagesize=LETTER, pageTemplates=[template], title="Guida operativa La Speziale", author="La Speziale")
    story = []
    list_number = 0
    pending_image = False

    for block in iter_blocks(source):
        if isinstance(block, DocxParagraph):
            if has_page_break(block):
                story.append(PageBreak())
                list_number = 0
                continue
            image = paragraph_image(block, source)
            if image:
                story.append(Spacer(1, 4))
                story.append(image)
                story.append(Spacer(1, 5))
                pending_image = True
                continue
            text = text_with_format(block).strip()
            if not text:
                continue
            style_name = block.style.name if block.style else "Normal"
            max_size = max((run.font.size.pt for run in block.runs if run.font.size), default=0)
            all_caps = block.text.strip().isupper() and len(block.text.strip()) < 70
            if style_name == "Heading 1":
                story.append(Paragraph(text, h1))
                list_number = 0
            elif style_name == "Heading 2":
                story.append(Paragraph(text, h2))
                list_number = 0
            elif style_name == "Heading 3":
                story.append(Paragraph(text, h3))
            elif style_name == "List Bullet":
                story.append(Paragraph(text, bullet, bulletText="•"))
            elif style_name == "List Number":
                list_number += 1
                story.append(Paragraph(text, number, bulletText=f"{list_number}."))
            elif max_size >= 24:
                story.append(Paragraph(text, title_style))
            elif max_size >= 12:
                story.append(Paragraph(text, subtitle_style))
            elif all_caps:
                story.append(Paragraph(text, kicker))
            elif pending_image and any(run.italic for run in block.runs):
                story.append(Paragraph(text, caption))
                pending_image = False
            else:
                story.append(Paragraph(text, body))
                pending_image = False
        else:
            data = []
            fills = []
            for r_idx, row in enumerate(block.rows):
                data_row = []
                fill_row = []
                for c_idx, cell in enumerate(row.cells):
                    text = "<br/>".join(escape(p.text) for p in cell.paragraphs if p.text.strip())
                    has_bold = any(run.bold for p in cell.paragraphs for run in p.runs)
                    data_row.append(Paragraph(text, cell_bold if has_bold else cell_style))
                    fill_row.append(cell_fill(cell))
                data.append(data_row)
                fills.append(fill_row)
            if not data:
                continue
            tbl = Table(data, colWidths=table_widths(block), repeatRows=1 if any(fills[0]) else 0, hAlign="LEFT")
            commands = [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D7DDD4")),
            ]
            for r_idx, row in enumerate(fills):
                for c_idx, fill in enumerate(row):
                    if fill:
                        commands.append(("BACKGROUND", (c_idx, r_idx), (c_idx, r_idx), fill))
                        if fill == SAGE:
                            commands.append(("TEXTCOLOR", (c_idx, r_idx), (c_idx, r_idx), WHITE))
            tbl.setStyle(TableStyle(commands))
            story.append(Spacer(1, 4))
            story.append(tbl)
            story.append(Spacer(1, 9))

    pdf.build(story)
    print(OUTPUT)


if __name__ == "__main__":
    build_pdf()
